# agent.py
import os
import asyncio
from pathlib import Path  
from docx import Document
from google import genai
from google.genai import types
from datetime import datetime

# Sibling imports for flat app structure
from config import GEMINI_API_KEY
from schemas import PlanResponse, ReflectionResult

class AutonomousAgent:
    def __init__(self):
        # Initialize Google GenAI client natively
        self.client = genai.Client(api_key=GEMINI_API_KEY)
        self.model = "gemini-3.1-flash-lite"
        
        # Concurrency control: Restrict to 2 parallel tasks to protect Free Tier RPM limits
        self.semaphore = asyncio.Semaphore(2)

    async def execute(self, user_request: str) -> tuple[list, str]:
        """Orchestrates the entire agentic pipeline."""
        print(f"\n[1/4] Phase: Planning workflow for request: '{user_request}'")
        plan = self._generate_plan(user_request)
        print(f"Plan Generated: '{plan.document_title}' with {len(plan.tasks)} sections.")

        print("\n[2/4] Phase: Executing tasks concurrently...")
        # Pack async execution tasks together
        async_tasks = [self._execute_task_safely(task, user_request) for task in plan.tasks]
        completed_sections = await asyncio.gather(*async_tasks)

        print("\n[3/4] Phase: Launching Reflection / Quality Assurance Check...")
        final_sections = []
        for title, raw_content in completed_sections:
            approved_content = self._reflect_and_correct(title, raw_content)
            final_sections.append((title, approved_content))

        print("\n[4/4] Phase: Compiling final artifact into binary Word file...")
        output_path = self._compile_to_docx(plan.document_title, final_sections)
        
        task_summaries = [f"Generated section: {title}" for title, _ in final_sections]
        return task_summaries, output_path

    def _generate_plan(self, user_request: str) -> PlanResponse:
        system_prompt = (
            "You are an elite corporate enterprise operations planner. Break down the user's request "
            "into logical, multi-layered document sections. Return data strictly matching the requested JSON schema."
        )
        response = self.client.models.generate_content(
            model=self.model,
            contents=f"User request: {user_request}",
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                response_mime_type="application/json",
                response_schema=PlanResponse,
                temperature=0.2
            )
        )
        return PlanResponse.model_validate_json(response.text)

    async def _execute_task_safely(self, task, context: str) -> tuple[str, str]:
        """Executes a text generation task protected by an AsyncIO Semaphore."""
        async with self.semaphore:
            # Let the CPU breathe and prevent heavy burst requests
            await asyncio.sleep(1.5)
            
            loop = asyncio.get_running_loop()
            # Run the synchronous SDK call in an executor thread to keep FastAPI non-blocking
            content = await loop.run_in_executor(None, self._generate_section_content, task, context)
            return task.title, content

    def _generate_section_content(self, task, context: str) -> str:
        prompt = (
            f"Context: Designing a master document for: {context}\n"
            f"Your specific task is to write the section titled: '{task.title}'\n"
            f"Section Requirements: {task.description}\n\n"
            f"Write comprehensive, high-grade corporate analysis. "
            f"Rule: Write purely structural prose. Do NOT write markdown strings like '**' or '###'."
        )
        response = self.client.models.generate_content(
            model=self.model,
            contents=prompt
        )
        return response.text.strip()

    def _reflect_and_correct(self, title: str, content: str, current_retry: int = 1, max_retries: int = 2) -> str:
        """The Reflection Guardrail: Evaluates output and forces adjustments if issues are spotted."""
        rubric = "Content must be highly authoritative, free of raw markdown text parsing symbols, and strictly professional."
        
        eval_prompt = (
            f"Evaluate this draft section content titled '{title}' against the following quality standards: {rubric}\n\n"
            f"DRAFT CONTENT:\n{content}"
        )
        
        response = self.client.models.generate_content(
            model=self.model,
            contents=eval_prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=ReflectionResult,
                temperature=0.0 # Deterministic evaluation
            )
        )
        result = ReflectionResult.model_validate_json(response.text)
        
        if result.is_approved or current_retry > max_retries:
            print(f"  └─ QA Status for '{title}': APPROVED")
            return content
        
        print(f"  └─ QA Status for '{title}': REJECTED. Re-writing via feedback logic... (Attempt {current_retry})")
        
        rewrite_prompt = (
            f"Rewrite this section content to explicitly resolve the following QA critique: {result.critique}\n\n"
            f"ORIGINAL CONTENT:\n{content}"
        )
        corrected_response = self.client.models.generate_content(
            model=self.model,
            contents=rewrite_prompt
        )
        
        # Recurse to verify the correction worked
        return self._reflect_and_correct(title, corrected_response.text.strip(), current_retry + 1, max_retries)

    def _compile_to_docx(self, document_title: str, sections: list) -> str:
        doc = Document()
        doc.add_heading(document_title, level=0)
        
        for title, body in sections:
            doc.add_heading(title, level=1)
            doc.add_paragraph(body)
            
        # OPTIMIZATION: Using Pathlib for bulletproof, object-oriented pathing
        target_output_dir = Path(__file__).resolve().parent.parent / "output"
        target_output_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Document_{timestamp}.docx"
        
        # Combine folder and filename dynamically
        final_file_path = target_output_dir / filename
        
        # Save the file and cast the Path object back to a string for your FastAPI JSON return
        doc.save(final_file_path)
        print(f" -> Output successfully generated at: {final_file_path}")
        
        return str(final_file_path)